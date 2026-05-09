from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE = BASE / "GOOGLE_CHALLENGE_SYSTEM_EXPLANATION.md"
OUTPUT = BASE / "WebA11y_Copilot_Explicacion_Sistema_Google.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(86, 96, 111)
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color="D7DBE2", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D7DBE2", size="4")
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F7F9FC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=9.2, color=RGBColor(45, 55, 72))
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("WebA11y Copilot - Documento de concurso")
    set_run_font(run, size=9, color=GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("WebA11y Copilot")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Documento de explicacion del sistema")
    set_run_font(run, size=15, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("Propuesta tecnica para concurso de Google AI")
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="B8C7DA", size="18")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="D7DBE2", size="6")
    rows = [
        ("Proyecto", "WebA11y Copilot"),
        ("Enfoque", "Accesibilidad web, WCAG 2.2 e inteligencia artificial"),
        ("Integracion Google AI", "Gemini 2.5 Flash como modelo operativo; Gemma 4 como fallback configurable"),
        ("Estado", "Probado localmente con 25 pruebas automatizadas aprobadas"),
    ]
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Inches(1.65)
        cells[1].width = Inches(4.85)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], TABLE_FILL)
        set_run_font(cells[0].paragraphs[0].add_run(label), size=10, color=NAVY, bold=True)
        set_run_font(cells[1].paragraphs[0].add_run(value), size=10, color=RGBColor(45, 55, 72))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Una herramienta practica para transformar hallazgos de accesibilidad en recomendaciones claras, "
        "accionables y utiles para equipos que construyen la web."
    )
    set_run_font(run, size=11, color=GRAY, italic=True)
    doc.add_page_break()


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D7DBE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY, bold=True)
    doc.add_paragraph()


def add_markdown_content(doc, markdown):
    in_code = False
    code_lines = []
    skip_title = True
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if skip_title and line.startswith("# "):
            skip_title = False
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[2:].strip())
            set_run_font(run, size=10.8, color=RGBColor(35, 45, 60))
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(re.sub(r"^\d+\. ", "", line))
            set_run_font(run, size=10.8, color=RGBColor(35, 45, 60))
            continue

        text = re.sub(r"`([^`]+)`", r"\1", line)
        if line.startswith("El proyecto fue probado") or line.startswith("Resultado actual:"):
            add_callout(doc, text)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.25
            run = p.add_run(text)
            set_run_font(run, size=11, color=RGBColor(32, 40, 52))


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_markdown_content(doc, markdown)
    doc.core_properties.title = "WebA11y Copilot - Documento de explicacion del sistema"
    doc.core_properties.subject = "Concurso Google AI"
    doc.core_properties.author = "Daniel Rivera Alpizar"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
